from __future__ import annotations

import datetime as dt
import hashlib
import json
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RC2_SRC = ROOT / "reference_candidate_rc2"
RC2_BUILD = ROOT / "build_rc2"
RC3_SRC = ROOT / "reference_candidate_rc3"
BUILD = ROOT / "build_rc3"
PACKAGE = BUILD / "package"
TEST = BUILD / "test"
ZIP = BUILD / "TRAJECTORY_FUNCTIONAL_REFERENCE_v0.9-rc3.zip"
VERSION = "0.9.0-rc3"
DISPLAY_VERSION = "v0.9-rc3"


def run(args: list[str], *, timeout: int = 150, stdout: Path | None = None, capture: bool = False) -> str:
    print("+", " ".join(args), flush=True)
    if stdout:
        with stdout.open("wb") as fh:
            subprocess.run(args, cwd=ROOT, check=True, timeout=timeout, stdout=fh)
        return ""
    if capture:
        proc = subprocess.run(args, cwd=ROOT, check=True, timeout=timeout, text=True, capture_output=True)
        if proc.stdout:
            print(proc.stdout.strip(), flush=True)
        if proc.stderr:
            print(proc.stderr.strip(), file=sys.stderr, flush=True)
        return proc.stdout.strip()
    subprocess.run(args, cwd=ROOT, check=True, timeout=timeout)
    return ""


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def write_text(name: str, value: str) -> None:
    (PACKAGE / name).write_text(textwrap.dedent(value).strip() + "\n", encoding="utf-8")


def build_rc2_base() -> None:
    shutil.rmtree(RC2_BUILD, ignore_errors=True)
    run([sys.executable, str((RC2_SRC / "patch_builder.py").relative_to(ROOT))])
    run([sys.executable, str((RC2_SRC / "build_rc2.py").relative_to(ROOT))], timeout=220)
    start = RC2_BUILD / "package" / "START.html"
    if not start.is_file() or start.stat().st_size == 0:
        raise RuntimeError("RC2 base START.html was not built")


def assemble_start() -> str:
    base = (RC2_BUILD / "package" / "START.html").read_text(encoding="utf-8")
    css = (RC3_SRC / "rc3_override.css").read_text(encoding="utf-8")
    js = (RC3_SRC / "rc3_override.js").read_text(encoding="utf-8")
    if "</head>" not in base or "</body>" not in base:
        raise RuntimeError("HTML injection anchors not found")
    source = base.replace("v0.9-rc2", DISPLAY_VERSION).replace("0.9.0-rc2", VERSION)
    source = source.replace("</head>", f"<style id='rc3-overrides'>\n{css}\n</style>\n</head>", 1)
    source = source.replace("</body>", f"<script id='rc3-runtime'>\n{js}\n</script>\n</body>", 1)
    (PACKAGE / "START.html").write_text(source, encoding="utf-8")
    return source


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=75, bottom=70, end=75) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def cell_text(cell, value: str, *, bold=False, color="172033", size=8.2, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if align is not None:
        paragraph.alignment = align
    run_ = paragraph.add_run(value)
    run_.bold = bold
    run_.font.name = "Aptos"
    run_._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run_.font.size = Pt(size)
    run_.font.color.rgb = RGBColor.from_string(color)


def checkbox_cell(cell, *, known: str | None = None) -> None:
    values = ["PASS", "FAIL", "N/A"]
    cell.text = ""
    for index, value in enumerate(values):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run_ = paragraph.add_run(("☒" if known == value else "☐") + " " + value)
        run_.font.name = "Aptos"
        run_._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        run_.font.size = Pt(8)
        run_.bold = known == value


def priority_cell(cell) -> None:
    cell.text = ""
    for index, value in enumerate(["P0", "P1", "P2", "—"]):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run_ = paragraph.add_run("☐ " + value)
        run_.font.name = "Aptos"
        run_._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        run_.font.size = Pt(8)


def build_protocol_docx() -> None:
    output = PACKAGE / "PROTOKOL_PRIEMKI_TRAJECTORY_REFERENCE_v0.9-rc3.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9)
    for name, size, color in (("Title", 18, "172033"), ("Heading 1", 14, "172033"), ("Heading 2", 11, "2563EB")):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Целевой протокол приёмки\nTRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    run_ = subtitle.add_run("Проверка изменения FOCUS-05 и короткая регрессия перед фиксацией v1.0")
    run_.font.color.rgb = RGBColor.from_string("536174")
    run_.font.size = Pt(9.5)

    info = doc.add_table(rows=1, cols=1)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    cell = info.cell(0, 0)
    set_cell_shading(cell, "EAF7EE")
    set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
    cell_text(
        cell,
        "ТЕХНИЧЕСКАЯ ПРОВЕРКА ВЫПОЛНЯЕТСЯ АВТОМАТИЧЕСКИ. Команде не нужно вручную проверять SHA-256, динамический каталог или JSON. В архив включены VERIFICATION_REPORT_RC3.md и AUTOTEST_RESULTS_RC3.json; публикация блокируется при техническом FAIL.",
        bold=True,
        color="166534",
        size=9.2,
    )

    doc.add_paragraph("Данные проверки", style="Heading 1")
    meta = doc.add_table(rows=3, cols=4)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.2, 8.8, 3.2, 8.8]
    values = [
        ("ФИО", "", "Роль", "☐ РБР   ☐ Методист   ☐ Технический специалист"),
        ("Дата", "", "Браузер / ОС", ""),
        ("Версия", DISPLAY_VERSION, "Итог", "☐ Принять как v1.0   ☐ Нужен rc4   ☐ Не принято"),
    ]
    for r, row_values in enumerate(values):
        for c, value in enumerate(row_values):
            target = meta.cell(r, c)
            set_cell_width(target, widths[c])
            set_cell_margins(target)
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            label = c % 2 == 0
            if label:
                set_cell_shading(target, "EAF0F8")
            cell_text(target, value, bold=label, color="344054" if label else "172033", size=8.6)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Как заполнять: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run("выполните действие, сравните результат с ожидаемым и отметьте один статус. При FAIL укажите P0/P1/P2 и фактическое поведение. Технические проверки из этого протокола исключены.")

    tests = [
        ("RUN-01", "Автономный запуск", "Распаковать ZIP и открыть START.html без интернета.", "Карта запускается без сервера, установки и загрузки дополнительных данных."),
        ("THEME-01", "Светлая тема", "Открыть карту, правую панель и same-semester блок в светлой теме.", "Текст, рамки, метки и кнопки читаемы; выделение параллельных дисциплин заметно."),
        ("THEME-02", "Тёмная тема", "Повторить проверку в тёмной теме.", "Все тексты и статусы контрастны; смысловые цвета компетенций сохраняются."),
        ("FOCUS-02", "Фокус с фильтром", "Выбрать компетенцию, затем дисциплину.", "Фильтр сохраняется; связи строятся только по активному набору компетенций."),
        ("FOCUS-03", "Связанные дисциплины других семестров", "Выбрать дисциплину с межсеместровыми совпадениями.", "Связанные карточки выделены и соединены цветными линиями; несвязанные приглушены."),
        ("FOCUS-05-RC3", "Совпадения в том же семестре", "Выбрать дисциплину, имеющую общую компетенцию с другой дисциплиной того же семестра.", "Параллельная дисциплина остаётся выделенной и имеет маркер «тот же семестр», но соединительная линия между ними отсутствует."),
        ("FOCUS-06", "Правая панель", "Просмотреть оба раздела связей и перейти по элементу списка.", "Есть разделы «Связанные дисциплины других семестров» и «Совпадения в текущем семестре»; общие компетенции и коэффициенты видны, переход работает."),
        ("SEMANTICS-01", "Пояснение семантики", "Прочитать пояснение в same-semester разделе.", "Интерфейс явно сообщает, что совпадение не является пререквизитом, постреквизитом или последовательностью изучения."),
        ("ISSUE-REG", "Регрессия проблем", "Открыть две проблемы подряд и вернуть одну в работу.", "Аккордеон раскрывается в месте клика; диагностика синхронна; статусы и комментарии работают."),
        ("RESET-01", "Снятие выделения", "Проверить «Снять фокус», «Снять компетенции», «Сбросить выделение» и Esc.", "Каждая команда снимает только заявленное состояние; повторная настройка фильтров не требуется."),
    ]

    doc.add_paragraph("Таблица целевой приёмки", style="Heading 1")
    headers = ["ID", "Что проверяем", "Что сделать", "Ожидаемый результат", "Отметка", "Приоритет при FAIL", "Комментарий / скрин"]
    column_widths = [1.65, 3.3, 5.1, 7.0, 2.0, 2.0, 5.0]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        target = table.cell(0, i)
        set_cell_width(target, column_widths[i])
        set_cell_shading(target, "1F4E78")
        set_cell_margins(target)
        cell_text(target, header, bold=True, color="FFFFFF", size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])

    for test_id, name, action, expected in tests:
        row = table.add_row()
        values = [test_id, name, action, expected]
        for index, value in enumerate(values):
            target = row.cells[index]
            set_cell_width(target, column_widths[index])
            set_cell_margins(target)
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell_text(target, value, bold=index in (0, 1), size=7.8)
        checkbox_cell(row.cells[4])
        priority_cell(row.cells[5])
        set_cell_width(row.cells[4], column_widths[4])
        set_cell_width(row.cells[5], column_widths[5])
        set_cell_width(row.cells[6], column_widths[6])
        set_cell_margins(row.cells[4])
        set_cell_margins(row.cells[5])
        set_cell_margins(row.cells[6])
        cell_text(row.cells[6], "", size=7.8)

    doc.add_page_break()
    doc.add_paragraph("Реестр дефектов", style="Heading 1")
    defect_headers = ["№", "ID теста", "Дефект", "Шаги", "Фактический результат", "Ожидаемый результат", "P0/P1/P2", "Скрин / видео"]
    defect_widths = [0.8, 1.6, 3.2, 4.5, 4.7, 4.7, 2.0, 3.5]
    defects = doc.add_table(rows=1, cols=len(defect_headers))
    defects.style = "Table Grid"
    defects.alignment = WD_TABLE_ALIGNMENT.CENTER
    defects.autofit = False
    for i, header in enumerate(defect_headers):
        target = defects.cell(0, i)
        set_cell_width(target, defect_widths[i])
        set_cell_shading(target, "7030A0")
        cell_text(target, header, bold=True, color="FFFFFF", size=7.7, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(defects.rows[0])
    for number in range(1, 6):
        row = defects.add_row()
        for index, value in enumerate([str(number), "", "", "", "", "", "", ""]):
            set_cell_width(row.cells[index], defect_widths[index])
            set_cell_margins(row.cells[index])
            cell_text(row.cells[index], value, size=7.7)

    doc.add_paragraph("Итоговое решение", style="Heading 1")
    decision = doc.add_table(rows=5, cols=2)
    decision.style = "Table Grid"
    decision.alignment = WD_TABLE_ALIGNMENT.CENTER
    decision_values = [
        ("Решение", "☐ Принять RC3 как функциональный эталон v1.0\n☐ Выпустить rc4 и перепроверить только FAIL\n☐ Выпустить rc4 и повторить полный протокол\n☐ Отклонить механику"),
        ("Обязательные исправления", "1. ______________________________________________\n2. ______________________________________________"),
        ("РБР: ФИО / подпись / дата", "__________________________________________________"),
        ("Методист: ФИО / подпись / дата", "__________________________________________________"),
        ("Технический специалист: ФИО / подпись / дата", "__________________________________________________"),
    ]
    for i, (label, value) in enumerate(decision_values):
        set_cell_width(decision.cell(i, 0), 6.3)
        set_cell_width(decision.cell(i, 1), 18.7)
        set_cell_shading(decision.cell(i, 0), "EAF0F8")
        set_cell_margins(decision.cell(i, 0), top=100, bottom=100)
        set_cell_margins(decision.cell(i, 1), top=100, bottom=100)
        cell_text(decision.cell(i, 0), label, bold=True, size=8.7)
        cell_text(decision.cell(i, 1), value, size=8.7)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ = footer.add_run("TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3 · целевая приёмка перед v1.0")
    run_.font.name = "Aptos"
    run_._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run_.font.size = Pt(8)
    run_.font.color.rgb = RGBColor.from_string("667085")

    doc.core_properties.title = "Протокол приёмки TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3"
    doc.core_properties.subject = "Целевая проверка FOCUS-05 и регрессия"
    doc.core_properties.author = "OpenAI"
    doc.save(output)

    reopened = Document(output)
    if len(reopened.tables) < 4:
        raise RuntimeError("DOCX QA failed: tables missing")
    if not any("FOCUS-05-RC3" in cell.text for table in reopened.tables for row in table.rows for cell in row.cells):
        raise RuntimeError("DOCX QA failed: FOCUS-05-RC3 missing")
    forbidden = {"COMP-07", "PACKAGE-01", "AUTO-01"}
    text = "\n".join(cell.text for table in reopened.tables for row in table.rows for cell in row.cells)
    if any(item in text for item in forbidden):
        raise RuntimeError("DOCX QA failed: technical manual checks remain")
    with zipfile.ZipFile(output) as archive:
        if archive.testzip() is not None or "word/document.xml" not in archive.namelist():
            raise RuntimeError("DOCX integrity failed")


def build_documents() -> None:
    input_contract = RC2_BUILD / "package" / "INPUT_FORMAT_CONTRACT.json"
    if input_contract.exists():
        shutil.copy2(input_contract, PACKAGE / input_contract.name)

    contract = {
        "package": "TRAJECTORY_FUNCTIONAL_REFERENCE",
        "version": VERSION,
        "candidate_only": True,
        "focus_same_semester": {
            "shared_active_competency": "highlight_card",
            "draw_line": False,
            "dim_unrelated": True,
            "right_panel_section": "Совпадения в текущем семестре",
            "semantic_warning": "Совпадение не означает пререквизит, постреквизит или последовательность изучения",
        },
        "focus_other_semesters": {
            "highlight_card": True,
            "one_line_per_shared_active_competency": True,
            "right_panel_section": "Связанные дисциплины других семестров",
        },
        "team_protocol": {
            "manual_technical_checks_removed": ["COMP-07", "PACKAGE-01", "AUTO-01"],
            "technical_verification": "automatic_and_blocking",
        },
    }
    (PACKAGE / "FUNCTIONAL_CONTRACT_RC3.json").write_text(json.dumps(contract, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    write_text("README.md", f'''
        # TRAJECTORY_FUNCTIONAL_REFERENCE {DISPLAY_VERSION}

        Малый кандидат после решения методиста по `FOCUS-05`.

        ## Запуск

        1. Полностью распакуйте ZIP.
        2. Откройте `START.html` из корня.
        3. Интернет, сервер и дополнительные файлы не требуются.

        ## Изменение RC3

        Если выбранная дисциплина имеет общую активную компетенцию с другой дисциплиной того же семестра:

        - параллельная дисциплина остаётся выделенной;
        - несвязанные дисциплины приглушаются;
        - линия внутри семестра не строится;
        - справа отображается отдельный раздел «Совпадения в текущем семестре»;
        - интерфейс поясняет, что совпадение не является пререквизитом или постреквизитом.

        Технические проверки выполняются автоматически и не включены в ручной протокол команды.
    ''')
    write_text("CHANGELOG_RC3.md", '''
        # Изменения v0.9-rc3

        1. Дисциплины одного семестра с общей активной компетенцией больше не приглушаются.
        2. Для них добавлены отдельная рамка и маркер «тот же семестр».
        3. Линии внутри одного семестра по-прежнему не строятся.
        4. Правая панель разделена на межсеместровые связи и совпадения текущего семестра.
        5. Добавлено явное пояснение семантики параллельного совпадения.
        6. Из командного протокола удалены COMP-07, PACKAGE-01 и AUTO-01.
        7. Контрольные суммы, динамический каталог и техническая целостность проверяются автоматически.
    ''')
    write_text("TEAM_ACCEPTANCE_SCENARIOS_RC3.md", '''
        # Целевая приёмка v0.9-rc3

        Основная проверка: `FOCUS-05-RC3`.

        - выбрать дисциплину с общей компетенцией у другой дисциплины того же семестра;
        - убедиться, что параллельная дисциплина выделена и не приглушена;
        - убедиться, что соединительной линии внутри семестра нет;
        - проверить маркер «тот же семестр»;
        - проверить отдельный блок справа и переход по нему;
        - кратко перепроверить межсеместровые линии, обе темы и аккордеон проблем.

        Технические SHA, JSON и динамический каталог вручную не проверяются.
    ''')
    build_protocol_docx()


def static_validate(source: str) -> list[str]:
    for path in PACKAGE.glob("*.json"):
        json.loads(path.read_text(encoding="utf-8"))
    if "<script src=" in source:
        raise RuntimeError("External script found")
    if re.search(r"\bfetch\s*\(", source):
        raise RuntimeError("fetch() call found")
    scripts = re.findall(r"<script(?:\s[^>]*)?>(.*?)</script>", source, flags=re.S | re.I)
    if len(scripts) < 3:
        raise RuntimeError("Expected base, RC2 and RC3 inline scripts")
    for index, script in enumerate(scripts):
        app = TEST / f"app-{index}.js"
        app.write_text(script, encoding="utf-8")
        run(["node", "--check", str(app.relative_to(ROOT))])
    required = [
        "sameSemesterMatches",
        "same-semester-related",
        "Совпадения в текущем семестре",
        "Связанные дисциплины других семестров",
        "RC3-SAME-SEMESTER-01",
        "data-same-semester-node",
    ]
    missing = [marker for marker in required if marker not in source]
    if missing:
        raise RuntimeError("Missing RC3 mechanics: " + ", ".join(missing))
    return ["all_json_parse", "no_external_scripts", "no_fetch", "inline_js_syntax", "required_rc3_markers", "docx_integrity"]


def inject_script(source: str, script: str) -> str:
    if "</body>" not in source:
        raise RuntimeError("Body close marker missing")
    return source.replace("</body>", f"<script>\n{script}\n</script>\n</body>", 1)


def build_fixtures(source: str) -> None:
    theme_script = lambda mode: f"""
    document.addEventListener('DOMContentLoaded',()=>setTimeout(()=>{{
      state.theme='{mode}';applyTheme();renderAll();
    }},300));
    """
    (TEST / "light.html").write_text(inject_script(source, theme_script("light")), encoding="utf-8")
    (TEST / "dark.html").write_text(inject_script(source, theme_script("dark")), encoding="utf-8")

    scenario = r'''
    document.addEventListener('DOMContentLoaded',()=>setTimeout(()=>{
      const checks={};
      const set=(name,value)=>{checks[name]=Boolean(value)};
      try{
        localStorage.clear();
        const api=window.__TRAJECTORY_REFERENCE_RC3__;
        const fixture=api.findSameSemesterFixture({requireCrossSemester:true});
        if(!fixture) throw new Error('same-semester fixture not found');
        state.roleId=fixture.role_id;
        state.activeCompIds=new Set();
        state.focusedNodeId=fixture.node.id;
        state.selectedIssueId=null;
        state.tab='details';
        renderAll();
        setTimeout(()=>{
          const sameCards=[...document.querySelectorAll('.discipline.same-semester-related')];
          set('same_semester_highlighted',sameCards.length>0);
          set('same_semester_not_dimmed',sameCards.every(card=>!card.classList.contains('focus-dim')));
          set('same_semester_badge',sameCards.some(card=>card.querySelector('.same-semester-badge')));
          set('same_semester_no_line',computeFocusEdges(fixture.node,new Set()).every(edge=>getNode(edge.from).semester!==getNode(edge.to).semester));
          set('cross_semester_lines_preserved',document.querySelectorAll('.line-competency').length>0);
          set('unrelated_dimmed',document.querySelectorAll('.discipline.focus-dim').length>0);
          set('right_panel_same_section',Boolean(document.querySelector('[data-rc3-same-semester-section="1"]')));
          set('right_panel_split',[...document.querySelectorAll('#panel h3')].some(h=>h.textContent.trim()==='Связанные дисциплины других семестров'));
          set('semantic_warning',document.querySelector('.rc3-same-semester-note')?.textContent.includes('не означает пререквизит'));
          const button=document.querySelector('[data-same-semester-node]');
          const target=button?.dataset.sameSemesterNode;
          if(button) button.click();
          set('same_semester_navigation',Boolean(target)&&state.focusedNodeId===target);
          focusIssue('MAP-001');
          set('issue_accordion_regression',Boolean(document.querySelector('.issue-card.expanded .issue-body')));
          set('rc3_builtin_self_test',document.documentElement.dataset.rc3SelfTest==='ok');
          const ok=Object.values(checks).every(Boolean);
          document.documentElement.dataset.rc3Scenario=ok?'ok':'failed';
          document.documentElement.dataset.rc3ScenarioResults=encodeURIComponent(JSON.stringify(checks));
          window.__RC3_SCENARIO_RESULTS__=checks;
          console.log('RC3_SCENARIO',JSON.stringify(checks));
        },1200);
      }catch(error){
        console.error(error);
        document.documentElement.dataset.rc3Scenario='failed';
        document.documentElement.dataset.rc3ScenarioError=error.message;
      }
    },1700));
    '''
    (TEST / "scenario.html").write_text(inject_script(source, scenario), encoding="utf-8")


def contrast_ratio(left: str, right: str) -> float:
    def luminance(value: str) -> float:
        value = value.lstrip("#")
        rgb = [int(value[i:i+2], 16) / 255 for i in (0, 2, 4)]
        rgb = [x / 12.92 if x <= 0.03928 else ((x + 0.055) / 1.055) ** 2.4 for x in rgb]
        return 0.2126 * rgb[0] + 0.7152 * rgb[1] + 0.0722 * rgb[2]
    l1, l2 = luminance(left), luminance(right)
    return (max(l1, l2) + 0.05) / (min(l1, l2) + 0.05)


def browser_validate() -> tuple[str, list[dict[str, object]], dict[str, float]]:
    chrome = shutil.which("google-chrome") or shutil.which("chromium")
    if not chrome:
        raise RuntimeError("Chrome/Chromium not found")
    version = run([chrome, "--version"], capture=True)
    results: list[dict[str, object]] = []
    for mode in ("light", "dark", "scenario"):
        url = f"file://{(TEST / f'{mode}.html').resolve()}"
        checks: dict[str, bool] = {}
        attempts = []
        text = ""
        ok = False
        for attempt in range(1, 4):
            profile = TEST / f"profile-{mode}-{attempt}"
            profile.mkdir()
            dom = TEST / f"{mode}-dom-{attempt}.html"
            run([
                chrome, "--headless=new", "--disable-gpu", "--no-sandbox",
                "--no-first-run", "--no-default-browser-check", "--disable-background-networking",
                "--disable-sync", "--disable-extensions", "--allow-file-access-from-files",
                f"--user-data-dir={profile.resolve()}", "--virtual-time-budget=22000", "--dump-dom", url,
            ], timeout=100, stdout=dom)
            text = dom.read_text(encoding="utf-8", errors="replace")
            checks = {
                "runtime_ready": 'data-runtime-ready="1"' in text,
                "base_self_test": 'data-self-test="ok"' in text,
                "rc2_ready": 'data-rc2-ready="1"' in text,
                "rc2_self_test": 'data-rc2-self-test="ok"' in text,
                "rc3_ready": 'data-rc3-ready="1"' in text,
                "rc3_self_test": 'data-rc3-self-test="ok"' in text,
            }
            if mode == "light":
                checks["theme_light"] = 'data-theme="light"' in text
            if mode == "dark":
                checks["theme_dark"] = 'data-theme="dark"' in text
            if mode == "scenario":
                checks["behavioral_scenario"] = 'data-rc3-scenario="ok"' in text
                match = re.search(r'data-rc3-scenario-results="([^"]*)"', text)
                error = re.search(r'data-rc3-scenario-error="([^"]*)"', text)
                print({"scenario_results": match.group(1) if match else None, "scenario_error": error.group(1) if error else None})
            ok = all(checks.values())
            attempts.append({"attempt": attempt, "ok": ok, "checks": dict(checks), "dom_size": len(text)})
            if ok:
                break
            print({"browser_retry": mode, "attempt": attempt, "checks": checks, "dom_size": len(text)})
        results.append({"id": f"chrome_{mode}", "ok": ok, "checks": checks, "attempts": attempts})
        if not ok:
            raise RuntimeError(f"Browser validation failed for {mode}: {checks}; attempts={attempts}")

    for mode in ("light", "dark"):
        run([
            chrome, "--headless=new", "--disable-gpu", "--no-sandbox", "--allow-file-access-from-files",
            "--window-size=1920,1080", f"--screenshot={str((TEST / f'{mode}.png').resolve())}",
            f"file://{(TEST / f'{mode}.html').resolve()}",
        ], timeout=90)

    contrast = {
        "light_text_panel": contrast_ratio("#172033", "#FFFFFF"),
        "light_muted_panel": contrast_ratio("#536174", "#FFFFFF"),
        "dark_text_panel": contrast_ratio("#F8FAFC", "#0F172A"),
        "dark_muted_panel": contrast_ratio("#CBD5E1", "#0F172A"),
        "dark_subtle_panel": contrast_ratio("#A8B3C7", "#0F172A"),
        "dark_danger": contrast_ratio("#FCA5A5", "#3F151B"),
        "dark_warning": contrast_ratio("#FCD34D", "#3A2A0A"),
        "dark_normalized": contrast_ratio("#E2E8F0", "#334155"),
    }
    if any(value < 4.5 for value in contrast.values()):
        raise RuntimeError(f"Contrast below 4.5: {contrast}")
    return version, results, contrast


def write_verification(chrome_version: str, static_results: list[str], browser_results: list[dict[str, object]], contrast: dict[str, float]) -> None:
    now = dt.datetime.now(dt.timezone.utc).isoformat()
    auto = {
        "package": "TRAJECTORY_FUNCTIONAL_REFERENCE",
        "version": VERSION,
        "created_at": now,
        "status": "PASS",
        "chrome_version": chrome_version,
        "static_checks": [{"id": item, "ok": True} for item in static_results],
        "browser_checks": browser_results,
        "contrast": contrast,
        "manual_technical_checks_required": False,
        "rc3_focus_rule": {
            "same_semester_highlight": True,
            "same_semester_line": False,
            "other_semesters_lines": True,
        },
    }
    (PACKAGE / "AUTOTEST_RESULTS_RC3.json").write_text(json.dumps(auto, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    scenario_attempts = next(item for item in browser_results if item["id"] == "chrome_scenario")["attempts"]
    report = f"""
    # Отчёт автоматической проверки {DISPLAY_VERSION}

    **Статус: PASS**

    Дата: `{now}`  
    Браузер: `{chrome_version}`

    ## Проверено автоматически

    - автономный `START.html`, отсутствие внешних скриптов и `fetch()`;
    - синтаксис всех inline JavaScript;
    - корректность JSON и DOCX;
    - светлая и тёмная темы, контраст не ниже 4.5:1 для проверенных пар;
    - выделение дисциплин того же семестра с общей активной компетенцией;
    - отсутствие соединительных линий внутри одного семестра;
    - сохранение межсеместровых цветных линий;
    - приглушение действительно несвязанных карточек;
    - отдельный раздел правой панели и переход по нему;
    - пояснение, что параллельное совпадение не является пререквизитом/постреквизитом;
    - регрессия аккордеона проблем;
    - целостность пакета и `START.html` в корне.

    ## Ручной протокол

    Технические пункты `COMP-07`, `PACKAGE-01` и `AUTO-01` удалены. Команда проверяет только пользовательскую механику.

    Попытки поведенческого сценария: `{json.dumps(scenario_attempts, ensure_ascii=False)}`.
    """
    write_text("VERIFICATION_REPORT_RC3.md", report)


def package_files() -> dict[str, object]:
    shutil.copy2(TEST / "light.png", PACKAGE / "PREVIEW_LIGHT.png")
    shutil.copy2(TEST / "dark.png", PACKAGE / "PREVIEW_DARK.png")
    start_hash = sha256(PACKAGE / "START.html")
    write_text("TECHNICAL_CHECKSUMS.txt", f'''
        START.html SHA-256:
        {start_hash}

        SHA-256 всего ZIP публикуется рядом с архивом в PACKAGE_SHA256.txt.
        Командная ручная проверка контрольной суммы не требуется.
    ''')

    manifest = {
        "package": "TRAJECTORY_FUNCTIONAL_REFERENCE",
        "version": VERSION,
        "created_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "start_file": "START.html",
        "autonomous_start": True,
        "candidate_only": True,
        "files": [],
    }
    manifest["files"] = [
        {"name": path.name, "size": path.stat().st_size, "sha256": sha256(path)}
        for path in sorted(PACKAGE.iterdir()) if path.is_file() and path.name != "PACKAGE_MANIFEST.json"
    ]
    (PACKAGE / "PACKAGE_MANIFEST.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(PACKAGE.iterdir()):
            if path.is_file():
                archive.write(path, path.name)
    with zipfile.ZipFile(ZIP) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("ZIP integrity failed")
        names = archive.namelist()
        if "START.html" not in names or any("/" in name.strip("/") or ".." in name for name in names):
            raise RuntimeError("Invalid ZIP root structure")
    result = {
        "zip": str(ZIP),
        "size": ZIP.stat().st_size,
        "sha256": sha256(ZIP),
        "start_sha256": start_hash,
        "files": len(names),
    }
    (BUILD / "PACKAGE_SHA256.txt").write_text(f"{result['sha256']}  {ZIP.name}\n", encoding="utf-8")
    (BUILD / "START_SHA256.txt").write_text(f"{start_hash}  START.html\n", encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False))
    return result


def main() -> None:
    shutil.rmtree(BUILD, ignore_errors=True)
    PACKAGE.mkdir(parents=True)
    TEST.mkdir(parents=True)
    build_rc2_base()
    source = assemble_start()
    build_documents()
    static_results = static_validate(source)
    build_fixtures(source)
    chrome_version, browser_results, contrast = browser_validate()
    write_verification(chrome_version, static_results, browser_results, contrast)
    package_files()


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"BUILD FAILED: {exc}", file=sys.stderr)
        raise
