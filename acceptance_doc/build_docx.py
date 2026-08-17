from pathlib import Path
from datetime import datetime, timezone
import hashlib, zipfile
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT=Path('acceptance_doc/out')
OUT.mkdir(parents=True,exist_ok=True)
DOCX=OUT/'PROTOKOL_PRIEMKI_TRAJECTORY_REFERENCE_v0.9-rc1.docx'

def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); shd=pr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); pr.append(shd)
    shd.set(qn('w:fill'),fill)

def margins(cell,top=70,start=70,bottom=70,end=70):
    pr=cell._tc.get_or_add_tcPr(); mar=pr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); pr.append(mar)
    for n,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el=mar.find(qn('w:'+n))
        if el is None: el=OxmlElement('w:'+n); mar.append(el)
        el.set(qn('w:w'),str(v)); el.set(qn('w:type'),'dxa')

def width(cell,cm):
    pr=cell._tc.get_or_add_tcPr(); el=pr.find(qn('w:tcW'))
    if el is None: el=OxmlElement('w:tcW'); pr.append(el)
    el.set(qn('w:w'),str(int(cm*567))); el.set(qn('w:type'),'dxa')

def repeat_header(row):
    pr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); pr.append(el)

def para(cell,text='',bold=False,color='172033',size=7.8,align=None):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=1.0
    if align is not None: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.name='Aptos'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    return p

def checkboxes(cell,items,checked=()):
    cell.text=''
    for i,item in enumerate(items):
        p=cell.paragraphs[0] if i==0 else cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=1
        r=p.add_run(('☒ ' if item in checked else '☐ ')+item); r.bold=item in checked; r.font.name='Aptos'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos'); r.font.size=Pt(7.8)

SECTIONS=[
('A. Запуск и автономность',[
('RUN-01','Автономный запуск','Распаковать ZIP и открыть START.html двойным щелчком при отключённом интернете.','Карта открывается без установки, сервера, внешних файлов и запроса выбрать CSV/JSON.'),
('RUN-02','Загрузка интерфейса','Дождаться встроенной самопроверки.','Видны БР, группы компетенций, семестры и статус «Самопроверка: OK»; нет пустого экрана.'),
('RUN-03','Повторный запуск','Обновить страницу и повторно открыть START.html.','Интерфейс остаётся работоспособным; встроенные данные не теряются.')]),
('B. Светлая и тёмная темы',[
('THEME-01','Светлая тема','Просмотреть верхнюю панель, карточки, фильтры, ошибки и поля ввода.','Основной, вторичный и служебный текст читаем; границы и статусы различимы.'),
('THEME-02','Тёмная тема','Выбрать тёмную тему и повторить просмотр всех панелей.','Нет неконтрастного текста; placeholder, disabled, ошибки и серые статусы читаемы.'),
('THEME-03','Цвет компетенции','Выбрать компетенцию, переключить тему и БР.','Компетенция сохраняет смысловой оттенок; в тёмной теме меняется только яркость.'),
('THEME-04','Сохранение темы','Выбрать тему и обновить страницу.','Выбранная тема восстанавливается.')]),
('C. Бизнес-роли',[
('ROLE-01','Переключение БР','Поочерёдно открыть Кибер и DevOps.','Показываются дисциплины выбранной БР; общая часть и ветвление различимы.'),
('ROLE-02','Сохранение фильтров','Выбрать несколько компетенций и переключить БР.','Активный набор сохраняется; метки и линии пересчитываются.'),
('ROLE-03','Общая реализация','Выбрать общую дисциплину и открыть детали.','Видны общая реализация и связанные БР; сущность не ведёт себя как независимый дубль.')]),
('D. Каталог и фильтры компетенций',[
('COMP-01','Без фильтра','Нажать «Снять все» и не выбирать дисциплину.','На карточках виден полный назначенный набор компетенций.'),
('COMP-02','Накопление групп','Включить целиком две разные группы.','Обе группы активны одновременно; вторая не заменяет первую.'),
('COMP-03','Выключение группы','При двух активных группах выключить одну.','Снимается только выбранная группа; другая остаётся.'),
('COMP-04','Индивидуальный выбор','Включить и выключить несколько отдельных компетенций.','Чекбоксы работают независимо и синхронизированы с групповыми кнопками.'),
('COMP-05','Метки карточек','Активировать одну компетенцию или небольшой набор.','На карточках остаются только выбранные назначенные компетенции.'),
('COMP-06','Нерелевантные карточки','Оставить компетенцию, которой нет в части дисциплин.','Такие карточки остаются на месте и приглушаются, а не исчезают.'),
('COMP-07','Динамический каталог','Проверить названия и число групп.','Интерфейс не зашит под 43 компетенции или фиксированные группы; каталог берётся из профиля.')]),
('E. Фокус дисциплины и связанные дисциплины',[
('FOCUS-01','Фокус без фильтра','Снять фильтры и выбрать дисциплину с общими компетенциями в других семестрах.','Показываются все её общие компетенции и связи.'),
('FOCUS-02','Фокус с фильтром','Выбрать компетенции, затем дисциплину.','Показываются только активные общие компетенции; клик не сбрасывает фильтр.'),
('FOCUS-03','Выделение связанных дисциплин','Выбрать дисциплину, имеющую общие компетенции с дисциплинами других семестров.','Выбранная карточка выделена отдельно; связанные карточки явно выделены, несвязанные приглушены; цветные линии и список общих компетенций справа позволяют понять связь.'),
('FOCUS-04','Цвет связи','Просмотреть несколько связанных карточек и линий.','Цвет метки, линии и расшифровки совпадает; одна компетенция не меняет цвет.'),
('FOCUS-05','Хронология','Выбрать дисциплину с общей компетенцией в том же семестре.','Линия внутри одного семестра не строится; связи идут слева направо.'),
('FOCUS-06','Панель связей','Перейти по связанной дисциплине из правой панели.','Показаны дисциплина, семестр, точные общие компетенции и коэффициенты; переход фокусирует карточку.')]),
('F. Линии и семантика',[
('LINE-01','Одна линия на компетенцию','Выбрать пару дисциплин с несколькими общими компетенциями.','Каждая компетенция имеет отдельную параллельную линию.'),
('LINE-02','Точки подключения','Просмотреть линии у карточек на разных уровнях.','Линии выходят/входят с внешних сторон карточек и не пересекают текст.'),
('LINE-03','Коэффициент','Сравнить одну компетенцию с разными коэффициентами.','Глубина различима толщиной и расшифрована в подсказке или панели.'),
('LINE-04','Soft Skills','Включить Soft Skills и выбрать дисциплину.','Soft Skills пунктирные и обозначают места оценивания, а не наследование.'),
('LINE-05','Продолжение','Найти дисциплину, идущую несколько семестров.','Продолжение показано отдельной серой линией.'),
('LINE-06','Глобальная траектория','Выбрать одну компетенцию без фокуса дисциплины.','Получается читаемая хронологическая трасса без полной сетки всех пар.')]),
('G. Проблемы и расхождения',[
('ISSUE-01','Единый реестр','Открыть «Проблемы и расхождения».','Видны четыре типа проблем с разными значками и типами.'),
('ISSUE-02','Навигация','Нажать на проблему конкретной БР/семестра.','Переключаются БР и семестр, выделяется контекст, открывается диагностика.'),
('ISSUE-03','Проблемы дисциплины','Выбрать дисциплину с проблемой.','Справа показаны те же issue_id, что в общем реестре; статусы синхронизированы.'),
('ISSUE-04','Полная диагностика','Развернуть проблему.','Видны правило, тип, приоритет, статус, текущее/ожидаемое, источники, узлы и действие.'),
('ISSUE-05','Типы и цвета','Сравнить четыре типа открытых проблем.','Методическая — красная (!), источники — янтарные (≠), сопоставление — фиолетовое (?), импорт — индиго (⚙); цвет не единственный признак.'),
('ISSUE-06','Закрытие','Попытаться закрыть проблему без комментария, затем с комментарием.','Без комментария блокируется; с комментарием сохраняется и становится серой.'),
('ISSUE-07','Статусы','Проверить «Не актуальна», «Нормализована» и возврат в работу.','Серые статусы различаются текстом и данными; возврат восстанавливает активное состояние.'),
('ISSUE-08','Автоисправление','Через редактор устранить динамическое нарушение.','После перевалидации проблема закрывается автоматически; при повторном нарушении открывается.'),
('ISSUE-09','Счётчики','Сравнить карточку с открытыми и нормализованными проблемами.','Активные и серые счётчики разделены; кандидаты не увеличивают число ошибок.')]),
('H. Редактирование',[
('EDIT-01','Текущая реализация','Изменить компетенцию в области «Только эта реализация».','Изменяется только выбранный экземпляр; область понятна до применения.'),
('EDIT-02','Общая реализация БР','Выбрать общую дисциплину и область нескольких БР.','До применения видно число объектов и БР; после применения изменение видно в обеих БР.'),
('EDIT-03','Каноническая дисциплина','Выбрать все реализации канонической дисциплины.','Массовое изменение выполняется только после явного выбора.'),
('EDIT-04','Коэффициент','Добавить/удалить компетенцию и изменить коэффициент.','Допустимы целые 1–4; неверное значение блокируется понятным сообщением.'),
('EDIT-05','Перевалидация','Применить изменение.','Обновляются метки, линии, проверки, реестр, счётчики и журнал.'),
('EDIT-06','Отмена','Нажать «Отменить изменение».','Последняя операция откатывается, интерфейс и проверки пересчитываются.')]),
('I. Сессия и экспорт',[
('SESSION-01','Сохранение сессии','Изменить фильтры, тему, статус и компетенции; сохранить сессию.','Создаётся JSON с данными, UI-состоянием, статусами и журналом.'),
('SESSION-02','Загрузка сессии','Обновить страницу и загрузить сохранённую сессию.','Полностью восстанавливаются БР, фильтры, объекты, тема, изменения и статусы.'),
('SESSION-03','Несовместимая сессия','Загрузить JSON другой версии/профиля.','Файл отклоняется и не применяется частично.'),
('EXPORT-01','Полный снимок','Экспортировать снимок.','Содержит всю текущую модель, профиль, каталог и проблемы.'),
('EXPORT-02','Журнал','Экспортировать журнал.','Содержит только операции до/после, время, область и комментарии.'),
('EXPORT-03','Одна проблема','Экспортировать выбранную проблему.','Содержит одну диагностику, источники, контекст и историю.')]),
('J. Навигация',[
('NAV-01','Поиск','Найти дисциплину по части названия и коду.','Остаются релевантные дисциплины; переход работает.'),
('NAV-02','Нулевые часы','Выключить и включить показ 0 часов.','Нулевые узлы скрываются/возвращаются без поломки карты.'),
('NAV-03','Только проблемы','Включить режим «Только проблемы».','Остаются дисциплины с активными проблемами; карта и реестр согласованы.'),
('NAV-04','FIT','Нажать FIT при восьми семестрах.','Карта реально масштабируется по ширине, а не только прокручивается в начало.'),
('NAV-05','Технические P2','Включить технические P2.','Они появляются отдельно и не смешиваются с методическими нарушениями.')])]

doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width,sec.page_height=sec.page_height,sec.page_width; sec.top_margin=Cm(1.1);sec.bottom_margin=Cm(1.1);sec.left_margin=Cm(1.0);sec.right_margin=Cm(1.0)
for st in doc.styles:
    if hasattr(st,'font'): st.font.name='Aptos'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos') if st._element.rPr is not None else None
styles=doc.styles; styles['Normal'].font.size=Pt(9); styles['Title'].font.size=Pt(18); styles['Title'].font.bold=True; styles['Heading 1'].font.size=Pt(14); styles['Heading 1'].font.bold=True
p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Протокол командной приёмки\nфункционального эталона визуализации траектории обучения')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Кандидат: TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc1 · для РБР, методиста и технического специалиста');r.font.size=Pt(9);r.font.color.rgb=RGBColor.from_string('536174')
warn=doc.add_table(rows=1,cols=1);warn.style='Table Grid';c=warn.cell(0,0);shade(c,'FDE9E7');margins(c,110,130,110,130);p=para(c,'ИЗВЕСТНЫЙ БЛОКИРУЮЩИЙ ДЕФЕКТ RC1',True,'991B1B',10);p2=c.add_paragraph('При выборе дисциплины недостаточно выделяются другие дисциплины, связанные с ней общими компетенциями. Строка FOCUS-03 заранее отмечена как FAIL / P0. До фиксации v1.0 необходимо восстановить выделение связанных карточек, приглушение несвязанных, цветные линии и расшифровку связей справа.');p2.paragraph_format.space_after=Pt(0)
doc.add_paragraph('Данные проверки',style='Heading 1')
meta=doc.add_table(rows=4,cols=4);meta.style='Table Grid';meta.alignment=WD_TABLE_ALIGNMENT.CENTER
meta_rows=[('ФИО специалиста','','Роль','☐ РБР   ☐ Методист   ☐ Технический специалист   ☐ Другое'),('Дата','','ОС и браузер',''),('Версия','TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc1','Бизнес-роль','☐ Кибер   ☐ DevOps   ☐ Обе'),('Итог','☐ Принято   ☐ Требуется rc2   ☐ Не принято','Дефекты','P0: ___   P1: ___   P2: ___')]
for i,row in enumerate(meta_rows):
    for j,val in enumerate(row):
        para(meta.cell(i,j),val,j%2==0,'344054' if j%2==0 else '172033',8.6); margins(meta.cell(i,j));
        if j%2==0: shade(meta.cell(i,j),'EAF0F8')
doc.add_paragraph('Как заполнять',style='Heading 1')
p=doc.add_paragraph('Выполните действие, сравните с ожидаемым результатом и отметьте один статус: PASS — работает точно; FAIL — есть расхождение; N/A — неприменимо. При FAIL обязательно укажите P0/P1/P2 и фактическое поведение. P0 блокирует фиксацию v1.0.');p.paragraph_format.space_after=Pt(7)
doc.add_paragraph('Основная таблица приёмки',style='Heading 1')
headers=['ID','Что проверяем','Что сделать','Ожидаемый результат','Отметка','Приоритет при FAIL','Фактический результат / комментарий']; widths=[1.55,3.05,5.05,6.45,2.05,2.0,5.45]
t=doc.add_table(rows=1,cols=7);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
for i,h in enumerate(headers): width(t.cell(0,i),widths[i]);shade(t.cell(0,i),'1F4E78');para(t.cell(0,i),h,True,'FFFFFF',8,WD_ALIGN_PARAGRAPH.CENTER);margins(t.cell(0,i));t.cell(0,i).vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
repeat_header(t.rows[0]); colors=['DCE6F1','E2F0D9','FFF2CC','E4DFEC','F4CCCC','D9EAF7','FCE4D6','E2F0D9','DDEBF7','EDEDED']; count=0
for sidx,(title,rows) in enumerate(SECTIONS):
    row=t.add_row();merged=row.cells[0]
    for x in row.cells[1:]: merged=merged.merge(x)
    shade(row.cells[0],colors[sidx]);para(row.cells[0],title,True,'172033',9.1);margins(row.cells[0])
    for tid,name,action,expected in rows:
        row=t.add_row();count+=1
        for j,val in enumerate([tid,name,action,expected]): width(row.cells[j],widths[j]);para(row.cells[j],val,j<2,'172033',7.7 if j>1 else 7.9);margins(row.cells[j]);row.cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
        known=tid=='FOCUS-03'
        checkboxes(row.cells[4],['PASS','FAIL','N/A'],['FAIL'] if known else [])
        checkboxes(row.cells[5],['P0','P1','P2','—'],['P0'] if known else [])
        width(row.cells[4],widths[4]);width(row.cells[5],widths[5]);width(row.cells[6],widths[6]);margins(row.cells[4]);margins(row.cells[5]);margins(row.cells[6])
        para(row.cells[6],'Известный дефект rc1: связанные карточки не получают достаточного выделения. Повторить тест после исправления в rc2.' if known else '',known,'991B1B' if known else '172033',7.7)
        if known:
            for c in row.cells: shade(c,'FDE9E7')
doc.add_page_break();doc.add_paragraph('Реестр обнаруженных дефектов',style='Heading 1');doc.add_paragraph('Заполняется для каждой строки FAIL. Для P0 приложите скриншот/видео и точные шаги воспроизведения.')
dh=['№','ID теста','Название','Шаги воспроизведения','Фактический результат','Ожидаемый результат','Приоритет / статус','Ссылка на скрин / видео'];dw=[.8,1.5,3.2,4.5,4.7,4.7,2.0,3.3]
dt=doc.add_table(rows=1,cols=8);dt.style='Table Grid';dt.alignment=WD_TABLE_ALIGNMENT.CENTER;dt.autofit=False
for i,h in enumerate(dh):width(dt.cell(0,i),dw[i]);shade(dt.cell(0,i),'7030A0');para(dt.cell(0,i),h,True,'FFFFFF',7.7,WD_ALIGN_PARAGRAPH.CENTER);margins(dt.cell(0,i));repeat_header(dt.rows[0])
known=['1','FOCUS-03','Не выделяются связанные дисциплины','1) Снять фильтры.\n2) Выбрать дисциплину с общими компетенциями.\n3) Повторить с активной компетенцией.','Выбранная карточка выделяется, но связанные карточки недостаточно различимы; связь трудно прочитать.','Связанные карточки выделены, несвязанные приглушены, линии окрашены по компетенциям, справа указан точный общий набор.','P0 · открыта','']
for r in range(1,7):
    row=dt.add_row(); vals=known if r==1 else [str(r),'','','','','','','']
    for i,v in enumerate(vals):width(row.cells[i],dw[i]);para(row.cells[i],v,r==1 and i in [1,2,6],'991B1B' if r==1 else '172033',7.5);margins(row.cells[i]);shade(row.cells[i],'FDE9E7') if r==1 else None
doc.add_paragraph('Итоговое решение команды',style='Heading 1')
final=doc.add_table(rows=6,cols=2);final.style='Table Grid';final.alignment=WD_TABLE_ALIGNMENT.CENTER
fin=[('Решение','☐ Принять как v1.0\n☐ Выпустить rc2 и повторить FAIL/P0–P1\n☐ Выпустить rc2 и повторить полный протокол\n☐ Отклонить механику'),('Обязательные исправления','1. Восстановить выделение связанных дисциплин по общим компетенциям.\n2. __________________________________\n3. __________________________________'),('Улучшения после v1.0','________________________________________'),('РБР: ФИО / подпись / дата','________________________________________'),('Методист: ФИО / подпись / дата','________________________________________'),('Технический специалист: ФИО / подпись / дата','________________________________________')]
for i,(a,b) in enumerate(fin):width(final.cell(i,0),6.2);width(final.cell(i,1),18.8);shade(final.cell(i,0),'EAF0F8');para(final.cell(i,0),a,True,'172033',8.6);para(final.cell(i,1),b,False,'172033',8.6);margins(final.cell(i,0),100,100,100,100);margins(final.cell(i,1),100,100,100,100)
for sec in doc.sections:
    p=sec.footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc1 · протокол командной приёмки · кандидат, не финальный эталон');r.font.size=Pt(8);r.font.color.rgb=RGBColor.from_string('667085')
doc.core_properties.title='Протокол командной приёмки функционального эталона визуализации траектории';doc.core_properties.author='OpenAI';doc.core_properties.comments='FOCUS-03 заранее отмечен как известный P0.'
doc.save(DOCX)
# QA
check=Document(DOCX); assert len(check.tables)>=4; assert any('FOCUS-03' in c.text for tb in check.tables for rw in tb.rows for c in rw.cells)
with zipfile.ZipFile(DOCX) as z: assert z.testzip() is None and 'word/document.xml' in z.namelist()
sha=hashlib.sha256(DOCX.read_bytes()).hexdigest();(OUT/'PROTOKOL_SHA256.txt').write_text(f'{sha}  {DOCX.name}\n',encoding='utf-8')
print({'file':str(DOCX),'size':DOCX.stat().st_size,'tests':count,'sha256':sha})
