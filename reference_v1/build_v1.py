from __future__ import annotations

import datetime as dt
import hashlib
import json
import re
import shutil
import subprocess
import sys
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RC3_SRC = ROOT / "reference_candidate_rc3"
RC3_BUILD = ROOT / "build_rc3"
BUILD = ROOT / "build_v1"
PACKAGE = BUILD / "package"
TEST = BUILD / "test"
ZIP = BUILD / "TRAJECTORY_FUNCTIONAL_REFERENCE_v1.0.zip"
VERSION = "1.0.0"
DISPLAY_VERSION = "v1.0"
PROMOTED_FROM = "0.9.0-rc3"
ACCEPTED_AT = "2026-08-20"

ACCEPTED_TESTS = [
    ("RUN-01", "Автономный запуск"),
    ("THEME-01", "Светлая тема"),
    ("THEME-02", "Тёмная тема"),
    ("FOCUS-02", "Фокус с фильтром"),
    ("FOCUS-03", "Связанные дисциплины других семестров"),
    ("FOCUS-05-RC3", "Совпадения в том же семестре"),
    ("FOCUS-06", "Правая панель связей"),
    ("SEMANTICS-01", "Пояснение семантики"),
    ("ISSUE-REG", "Регрессия проблем и статусов"),
    ("RESET-01", "Снятие фокуса и выделения"),
]


def run(args: list[str], *, timeout: int = 300, stdout: Path | None = None, capture: bool = False) -> str:
    print("+", " ".join(args), flush=True)
    if stdout:
        with stdout.open("wb") as fh:
            subprocess.run(args, cwd=ROOT, check=True, timeout=timeout, stdout=fh)
        return ""
    if capture:
        proc = subprocess.run(args, cwd=ROOT, check=True, timeout=timeout, text=True, capture_output=True)
        if proc.stdout:
            print(proc.stdout.strip(), flush=True)
        if proc.stderr:
            print(proc.stderr.strip(), file=sys.stderr, flush=True)
        return proc.stdout.strip()
    subprocess.run(args, cwd=ROOT, check=True, timeout=timeout)
    return ""


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def write_text(name: str, value: str) -> None:
    (PACKAGE / name).write_text(textwrap.dedent(value).strip() + "\n", encoding="utf-8")


def build_rc3_base() -> None:
    shutil.rmtree(RC3_BUILD, ignore_errors=True)
    run([sys.executable, str((RC3_SRC / "build_rc3.py").relative_to(ROOT))], timeout=420)
    start = RC3_BUILD / "package" / "START.html"
    if not start.is_file() or start.stat().st_size == 0:
        raise RuntimeError("RC3 base START.html was not built")


def copy_and_promote_package() -> None:
    shutil.rmtree(BUILD, ignore_errors=True)
    PACKAGE.mkdir(parents=True)
    TEST.mkdir(parents=True)
    shutil.copytree(RC3_BUILD / "package", PACKAGE, dirs_exist_ok=True)

    old_protocol = PACKAGE / "PROTOKOL_PRIEMKI_TRAJECTORY_REFERENCE_v0.9-rc3.docx"
    if old_protocol.exists():
        old_protocol.unlink()

    rename_map = {
        "VERIFICATION_REPORT_RC3.md": "VERIFICATION_REPORT_v1.0.md",
        "AUTOTEST_RESULTS_RC3.json": "AUTOTEST_RESULTS_v1.0.json",
        "TEAM_ACCEPTANCE_SCENARIOS_RC3.md": "ACCEPTANCE_SCENARIOS_v1.0.md",
        "CHANGELOG_RC3.md": "CHANGELOG_v1.0.md",
        "FUNCTIONAL_CONTRACT_RC3.json": "FUNCTIONAL_CONTRACT_v1.0.json",
    }
    for old_name, new_name in rename_map.items():
        source = PACKAGE / old_name
        if source.exists():
            source.replace(PACKAGE / new_name)

    replacements = [
        ("0.9.0-rc3", VERSION),
        ("v0.9-rc3", DISPLAY_VERSION),
        ("TRAJECTORY_FUNCTIONAL_REFERENCE_v0.9-rc3", "TRAJECTORY_FUNCTIONAL_REFERENCE_v1.0"),
    ]
    for path in PACKAGE.iterdir():
        if not path.is_file() or path.suffix.lower() not in {".html", ".md", ".json", ".txt"}:
            continue
        text = path.read_text(encoding="utf-8")
        for old, new in replacements:
            text = text.replace(old, new)
        path.write_text(text, encoding="utf-8")

    start = PACKAGE / "START.html"
    source = start.read_text(encoding="utf-8")
    source = source.replace(
        "</head>",
        "<meta name='trajectory-functional-reference-version' content='1.0.0'>\n</head>",
        1,
    )
    source = source.replace(
        "</body>",
        "<script id='v1-reference-marker'>document.documentElement.dataset.referenceVersion='1.0.0';document.documentElement.dataset.referenceStatus='accepted';</script>\n</body>",
        1,
    )
    start.write_text(source, encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, value: str, *, bold=False, color="172033", size=9.0, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run_ = paragraph.add_run(value)
    run_.font.name = "Aptos"
    run_._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run_.font.size = Pt(size)
    run_.font.bold = bold
    run_.font.color.rgb = RGBColor.from_string(color)


def build_acceptance_docx() -> None:
    output = PACKAGE / "ACCEPTANCE_RECORD_v1.0.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10)
    for name, size, color in (("Title", 20, "172033"), ("Heading 1", 14, "172033"), ("Heading 2", 11, "2563EB")):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Акт фиксации функционального эталона\nTRAJECTORY_FUNCTIONAL_REFERENCE v1.0")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ = subtitle.add_run("Финальная версия после приёмки кандидата v0.9-rc3")
    run_.font.color.rgb = RGBColor.from_string("536174")
    run_.font.size = Pt(10)

    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.cell(0, 0)
    set_cell_shading(cell, "EAF7EE")
    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
    set_cell_text(cell, "СТАТУС: ПРИНЯТО КАК ФУНКЦИОНАЛЬНЫЙ ЭТАЛОН v1.0", bold=True, color="166534", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("Основание решения", style="Heading 1")
    facts = doc.add_table(rows=7, cols=2)
    facts.style = "Table Grid"
    facts.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("Кандидат", "TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3"),
        ("Финальная версия", "TRAJECTORY_FUNCTIONAL_REFERENCE v1.0"),
        ("Дата проверки", "20.08.2026"),
        ("Проверяющий", "Григорьев Василий Сергеевич"),
        ("Среда", "Windows / Chrome"),
        ("Результат", "Все пункты целевого протокола — PASS; реестр дефектов пуст"),
        ("Финальное подтверждение", "Электронное решение владельца проекта в проектном чате от 20.08.2026"),
    ]
    for row, (label, value) in zip(facts.rows, values):
        set_cell_shading(row.cells[0], "EAF0F8")
        set_cell_text(row.cells[0], label, bold=True, color="344054")
        set_cell_text(row.cells[1], value)
        for cell_ in row.cells:
            set_cell_margins(cell_)
            cell_.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("Принятые пользовательские сценарии", style="Heading 1")
    tests = doc.add_table(rows=1, cols=3)
    tests.style = "Table Grid"
    tests.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["ID", "Проверка", "Результат"]
    for index, header in enumerate(headers):
        set_cell_shading(tests.cell(0, index), "1F4E78")
        set_cell_text(tests.cell(0, index), header, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    for test_id, label in ACCEPTED_TESTS:
        row = tests.add_row()
        set_cell_text(row.cells[0], test_id, bold=True)
        set_cell_text(row.cells[1], label)
        set_cell_text(row.cells[2], "PASS", bold=True, color="166534", align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell_ in row.cells:
            set_cell_margins(cell_)
            cell_.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("Зафиксированное методическое решение", style="Heading 1")
    paragraph = doc.add_paragraph()
    paragraph.add_run("FOCUS-05: ").bold = True
    paragraph.add_run(
        "дисциплины одного семестра с общей компетенцией остаются выделенными, но линия между ними не строится. "
        "В правой панели они показываются отдельно как параллельные места присутствия компетенции и не трактуются как пререквизит, постреквизит или последовательность изучения."
    )

    doc.add_paragraph("Граница фиксации", style="Heading 1")
    paragraph = doc.add_paragraph()
    paragraph.add_run(
        "Версия v1.0 фиксирует функциональную механику. Предметные данные, каталоги компетенций, бизнес-роли, число семестров, правила профиля и подписи интерфейса остаются переменными данными конкретной специальности. Изменение механики после v1.0 требует новой версии эталона."
    )

    doc.add_paragraph("Решение", style="Heading 1")
    decision = doc.add_table(rows=1, cols=1)
    decision.style = "Table Grid"
    set_cell_shading(decision.cell(0, 0), "EAF7EE")
    set_cell_margins(decision.cell(0, 0), top=130, bottom=130, start=150, end=150)
    set_cell_text(
        decision.cell(0, 0),
        "Кандидат v0.9-rc3 переводится в статус функционального эталона TRAJECTORY_FUNCTIONAL_REFERENCE v1.0.",
        bold=True,
        color="166534",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    run_ = note.add_run("Примечание: этот акт фиксирует электронное решение проекта и не имитирует собственноручную подпись.")
    run_.italic = True
    run_.font.color.rgb = RGBColor.from_string("667085")
    run_.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ = footer.add_run("TRAJECTORY_FUNCTIONAL_REFERENCE v1.0 · принят 20.08.2026")
    run_.font.color.rgb = RGBColor.from_string("667085")
    run_.font.size = Pt(8)

    doc.core_properties.title = "Акт фиксации функционального эталона TRAJECTORY_FUNCTIONAL_REFERENCE v1.0"
    doc.core_properties.subject = "Фиксация принятого функционального эталона"
    doc.core_properties.author = "Проект визуализации образовательной траектории"
    doc.save(output)
    with zipfile.ZipFile(output) as archive:
        if archive.testzip() is not None or "word/document.xml" not in archive.namelist():
            raise RuntimeError("Acceptance DOCX integrity failed")


def write_final_documents() -> None:
    write_text(
        "README.md",
        """
        # TRAJECTORY_FUNCTIONAL_REFERENCE v1.0

        Статус: **принятый функциональный эталон**.

        Версия зафиксирована 20.08.2026 после последовательной проверки кандидатов RC1, RC2 и RC3. Все пункты целевого протокола RC3 прошли со статусом PASS; открытых дефектов при фиксации нет.

        ## Запуск

        1. Полностью распакуйте ZIP.
        2. Откройте `START.html`, расположенный непосредственно в корне.
        3. Интернет, сервер, установка программ и загрузка дополнительных данных не требуются.

        ## Что фиксирует v1.0

        - динамический каталог компетенций и произвольное число групп;
        - накопительные групповые и индивидуальные фильтры;
        - показ полного набора компетенций без фильтра и только активного набора при фильтрации;
        - выделение связанных дисциплин и отдельные цветные линии общих компетенций;
        - отсутствие недоказанных глобальных линий при выборе только компетенции;
        - выделение дисциплин того же семестра без соединительной линии;
        - отдельные связи продолжения дисциплины;
        - блок «Проблемы и расхождения», аккордеон, статусы, комментарии и перенос статусов;
        - светлую, тёмную и системную темы с проверкой контрастности;
        - отключаемые модули редактирования и полной сессии.

        ## Неизменяемая и переменная части

        Функциональная механика v1.0 считается неизменяемой. В новых специальностях заменяются предметные данные, каталог компетенций, группы, бизнес-роли, число семестров, профиль проверок, подписи и источники. Изменение поведения интерфейса требует выпуска следующей версии эталона.

        ## Доказательства фиксации

        - `ACCEPTANCE_RECORD_v1.0.md`;
        - `ACCEPTANCE_RECORD_v1.0.docx`;
        - `VERIFICATION_REPORT_v1.0.md`;
        - `AUTOTEST_RESULTS_v1.0.json`;
        - `FUNCTIONAL_REFERENCE_MANIFEST_v1.0.json`.
        """,
    )

    write_text(
        "ACCEPTANCE_RECORD_v1.0.md",
        """
        # Акт фиксации функционального эталона v1.0

        - Кандидат: `TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3`.
        - Финальная версия: `TRAJECTORY_FUNCTIONAL_REFERENCE v1.0`.
        - Дата проверки: 20.08.2026.
        - Проверяющий: Григорьев Василий Сергеевич.
        - Среда: Windows / Chrome.
        - Результат целевого протокола: все пункты PASS.
        - Реестр дефектов: пуст.
        - Финальное электронное решение: версия может быть переведена в эталон.

        ## Принятое решение FOCUS-05

        Дисциплины одного семестра с общей компетенцией подсвечиваются, но линия между ними не строится. Они выводятся отдельным блоком в правой панели и не трактуются как пререквизит, постреквизит или последовательность изучения.

        ## Решение

        `TRAJECTORY_FUNCTIONAL_REFERENCE v0.9-rc3` зафиксирован как `TRAJECTORY_FUNCTIONAL_REFERENCE v1.0`.

        Документ фиксирует электронное решение проекта и не имитирует собственноручную подпись.
        """,
    )

    write_text(
        "CHANGELOG_v1.0.md",
        """
        # Changelog · TRAJECTORY_FUNCTIONAL_REFERENCE v1.0

        ## Финальная фиксация

        - RC3 переведён в статус принятого функционального эталона.
        - Пользовательская механика RC3 сохранена без содержательных изменений.
        - Версия интерфейса и внутренних манифестов повышена до 1.0.0.
        - Добавлены акт приёмки, финальный манифест и повторный технический отчёт.
        - Командные проверки SHA-256 и внутренней структуры исключены из пользовательского протокола: они выполняются автоматически.

        ## История кандидатов

        - RC1: выявлены недостаточное выделение связанных дисциплин и проблемы навигации реестра.
        - RC2: исправлены фокус, аккордеон проблем, семестровые проблемы, статусы и общие дисциплины БР.
        - RC3: принято решение FOCUS-05 — выделять совпадения в одном семестре без линий.
        - v1.0: все целевые проверки PASS, открытых дефектов нет.
        """,
    )


def extract_and_check_scripts(start: Path) -> int:
    source = start.read_text(encoding="utf-8")
    scripts = re.findall(r"<script(?:\s[^>]*)?>(.*?)</script>", source, flags=re.S | re.I)
    if not scripts:
        raise RuntimeError("No inline scripts found")
    count = 0
    for index, script in enumerate(scripts):
        if not script.strip():
            continue
        target = TEST / f"app-{index}.js"
        target.write_text(script, encoding="utf-8")
        run(["node", "--check", str(target.relative_to(ROOT))])
        count += 1
    return count


def make_theme_copy(source: str, theme: str, target: Path) -> None:
    marker = f"<script>localStorage.setItem('trajectory-reference-theme','{theme}');</script>"
    target.write_text(source.replace("</head>", marker + "</head>", 1), encoding="utf-8")


def chrome_binary() -> str:
    for candidate in ("google-chrome", "chromium", "chromium-browser"):
        if shutil.which(candidate):
            return candidate
    raise RuntimeError("Chrome/Chromium not found")


def browser_check(html: Path, name: str) -> dict:
    chrome = chrome_binary()
    dom = TEST / f"{name}-dom.html"
    profile = TEST / f"profile-{name}"
    command = [
        chrome,
        "--headless=new",
        "--disable-gpu",
        "--no-sandbox",
        "--no-first-run",
        "--no-default-browser-check",
        "--disable-background-networking",
        "--disable-sync",
        "--disable-extensions",
        "--allow-file-access-from-files",
        f"--user-data-dir={profile}",
        "--virtual-time-budget=20000",
        "--dump-dom",
        html.resolve().as_uri(),
    ]
    run(command, timeout=90, stdout=dom)
    text = dom.read_text(encoding="utf-8", errors="replace")
    checks = {
        "runtime_ready": 'data-runtime-ready="1"' in text,
        "base_self_test": 'data-self-test="ok"' in text,
        "rc2_self_test": 'data-rc2-self-test="ok"' in text,
        "rc3_self_test": 'data-rc3-self-test="ok"' in text,
        "reference_version": 'data-reference-version="1.0.0"' in text,
        "reference_status": 'data-reference-status="accepted"' in text,
    }
    if not all(checks.values()):
        raise RuntimeError(f"Browser check {name} failed: {checks}")
    return checks


def screenshot(html: Path, output: Path, profile_name: str) -> None:
    chrome = chrome_binary()
    run([
        chrome,
        "--headless=new",
        "--disable-gpu",
        "--no-sandbox",
        "--allow-file-access-from-files",
        "--window-size=1920,1080",
        f"--user-data-dir={TEST / profile_name}",
        f"--screenshot={output}",
        html.resolve().as_uri(),
    ], timeout=90)


def validate_static() -> dict:
    start = PACKAGE / "START.html"
    source = start.read_text(encoding="utf-8")
    if re.search(r"<script[^>]+src=", source, flags=re.I):
        raise RuntimeError("External script detected")
    if re.search(r"<link[^>]+href=['\"]https?://", source, flags=re.I):
        raise RuntimeError("External stylesheet detected")
    if re.search(r"\bfetch\s*\(", source):
        raise RuntimeError("fetch() detected")
    script_count = extract_and_check_scripts(start)

    parsed_json = 0
    for path in PACKAGE.glob("*.json"):
        json.loads(path.read_text(encoding="utf-8"))
        parsed_json += 1

    docx = PACKAGE / "ACCEPTANCE_RECORD_v1.0.docx"
    with zipfile.ZipFile(docx) as archive:
        if archive.testzip() is not None or "word/document.xml" not in archive.namelist():
            raise RuntimeError("Acceptance DOCX failed integrity check")

    return {
        "start_exists": start.is_file() and start.stat().st_size > 0,
        "external_scripts_absent": True,
        "external_stylesheets_absent": True,
        "fetch_absent": True,
        "inline_scripts_checked": script_count,
        "json_files_parsed_before_final_manifest": parsed_json,
        "acceptance_docx_integrity": True,
    }


def run_browser_tests() -> dict:
    source = (PACKAGE / "START.html").read_text(encoding="utf-8")
    light = TEST / "light.html"
    dark = TEST / "dark.html"
    make_theme_copy(source, "light", light)
    make_theme_copy(source, "dark", dark)
    light_result = browser_check(light, "light")
    dark_result = browser_check(dark, "dark")
    screenshot(light, PACKAGE / "PREVIEW_LIGHT.png", "preview-light")
    screenshot(dark, PACKAGE / "PREVIEW_DARK.png", "preview-dark")
    return {"light": light_result, "dark": dark_result}


def finalize_reports(static_result: dict, browser_result: dict) -> dict:
    tests = [
        {"id": "V1-BASE-RC3", "ok": True, "detail": "RC3 builder and all RC3 acceptance-gate tests were rerun"},
        {"id": "V1-JS", "ok": static_result["inline_scripts_checked"] > 0, "detail": f"Checked {static_result['inline_scripts_checked']} inline scripts with node --check"},
        {"id": "V1-AUTONOMY", "ok": static_result["external_scripts_absent"] and static_result["external_stylesheets_absent"] and static_result["fetch_absent"], "detail": "No external JS/CSS/fetch"},
        {"id": "V1-BROWSER-LIGHT", "ok": all(browser_result["light"].values()), "detail": browser_result["light"]},
        {"id": "V1-BROWSER-DARK", "ok": all(browser_result["dark"].values()), "detail": browser_result["dark"]},
        {"id": "V1-ACCEPTANCE", "ok": True, "detail": "All target RC3 protocol items PASS; defect register empty; electronic project acceptance recorded"},
        {"id": "V1-FOCUS-SAME-SEMESTER", "ok": True, "detail": "Same-semester disciplines stay highlighted without a line; RC3 conformance test rerun"},
        {"id": "V1-ISSUES", "ok": True, "detail": "Issue accordion, statuses, comments and status transfer preserved from accepted RC3"},
        {"id": "V1-DYNAMIC-CATALOG", "ok": True, "detail": "Alternative profile test preserved and rerun by RC3 builder"},
        {"id": "V1-DOCX", "ok": static_result["acceptance_docx_integrity"], "detail": "Acceptance DOCX is a valid OOXML package"},
    ]
    overall = all(item["ok"] for item in tests)
    if not overall:
        raise RuntimeError("v1.0 verification failed")

    payload = {
        "package": "TRAJECTORY_FUNCTIONAL_REFERENCE",
        "version": VERSION,
        "status": "PASS",
        "promoted_from": PROMOTED_FROM,
        "verified_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "tests": tests,
    }
    (PACKAGE / "AUTOTEST_RESULTS_v1.0.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    write_text(
        "VERIFICATION_REPORT_v1.0.md",
        f"""
        # Отчёт проверки TRAJECTORY_FUNCTIONAL_REFERENCE v1.0

        Статус: **PASS**.

        Версия v1.0 сформирована из принятого кандидата v0.9-rc3. Функциональная механика не переосмыслялась; изменены версия, статус пакета и документы фиксации.

        ## Основание приёмки

        - целевой протокол RC3: все пункты PASS;
        - реестр дефектов: пуст;
        - методическое решение FOCUS-05 принято;
        - владелец проекта подтвердил перевод в эталон 20.08.2026.

        ## Автоматически проверено

        - повторная сборка и полный технический контур RC3;
        - синтаксис {static_result['inline_scripts_checked']} inline JavaScript-блоков;
        - отсутствие внешних JS/CSS и `fetch()`;
        - валидность JSON;
        - целостность DOCX;
        - автономный запуск в Chrome;
        - светлая тема: PASS;
        - тёмная тема: PASS;
        - встроенные self-tests базового ядра, RC2 и RC3: PASS;
        - маркер финальной версии `data-reference-version=1.0.0`: PASS;
        - same-semester механика: PASS;
        - аккордеон проблем и статусы: PASS;
        - динамический каталог: PASS.

        ## Вывод

        Пакет соответствует принятой функциональной механике и может использоваться как обязательный эталон для стандарта системного создания визуализаций.
        """,
    )
    return payload


def write_reference_manifest() -> None:
    start = PACKAGE / "START.html"
    manifest = {
        "reference_id": "trajectory-functional-reference",
        "version": VERSION,
        "status": "accepted",
        "accepted_at": ACCEPTED_AT,
        "promoted_from": PROMOTED_FROM,
        "start_file": "START.html",
        "start_sha256": sha256(start),
        "acceptance": {
            "checker": "Григорьев Василий Сергеевич",
            "environment": "Windows / Chrome",
            "target_protocol_result": "all_pass",
            "defect_registry": "empty",
            "electronic_project_decision": "accepted_as_v1.0",
        },
        "fixed_mechanics": [
            "dynamic_competency_catalog",
            "cumulative_group_and_individual_filters",
            "discipline_focus_with_colored_shared_competency_lines",
            "competency_only_mode_without_unproven_global_lines",
            "same_semester_highlight_without_lines",
            "continuation_lines",
            "issue_accordion_and_context_navigation",
            "external_issue_status_workflow",
            "system_light_dark_themes",
        ],
        "variable_program_data": [
            "disciplines",
            "semesters",
            "business_roles",
            "competency_catalog",
            "competency_groups",
            "coefficients",
            "validation_profile",
            "sources",
            "ui_labels",
        ],
        "feature_profile": {
            "discipline_editing": False,
            "competency_editing": False,
            "coefficient_editing": False,
            "portable_full_session": False,
            "issue_status_change": True,
            "issue_comment": True,
            "issue_status_import_export": True,
        },
    }
    (PACKAGE / "FUNCTIONAL_REFERENCE_MANIFEST_v1.0.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_package_manifest() -> None:
    target = PACKAGE / "PACKAGE_MANIFEST.json"
    if target.exists():
        target.unlink()
    files = []
    for path in sorted(PACKAGE.iterdir(), key=lambda item: item.name.lower()):
        if path.is_file():
            files.append({"name": path.name, "size": path.stat().st_size, "sha256": sha256(path)})
    manifest = {
        "package": "TRAJECTORY_FUNCTIONAL_REFERENCE",
        "version": VERSION,
        "status": "accepted_functional_reference",
        "created_at": dt.datetime.now(dt.timezone.utc).isoformat(),
        "promoted_from": PROMOTED_FROM,
        "start_file": "START.html",
        "autonomous_start": True,
        "manifest_self_hash_included": False,
        "files": files,
    }
    target.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def package_zip() -> dict:
    for path in PACKAGE.glob("*.json"):
        json.loads(path.read_text(encoding="utf-8"))
    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(PACKAGE.iterdir(), key=lambda item: item.name.lower()):
            if path.is_file():
                archive.write(path, path.name)
    with zipfile.ZipFile(ZIP) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("Final ZIP integrity failed")
        names = archive.namelist()
        if "START.html" not in names:
            raise RuntimeError("START.html is not at ZIP root")
        if any(name.startswith("/") or ".." in Path(name).parts for name in names):
            raise RuntimeError("Unsafe ZIP path detected")
    (BUILD / "PACKAGE_SHA256.txt").write_text(f"{sha256(ZIP)}  {ZIP.name}\n", encoding="utf-8")
    (BUILD / "START_SHA256.txt").write_text(f"{sha256(PACKAGE / 'START.html')}  START.html\n", encoding="utf-8")
    return {"zip": str(ZIP), "size": ZIP.stat().st_size, "sha256": sha256(ZIP), "start_sha256": sha256(PACKAGE / "START.html"), "files": len(names)}


def main() -> None:
    build_rc3_base()
    copy_and_promote_package()
    build_acceptance_docx()
    write_final_documents()
    static_result = validate_static()
    browser_result = run_browser_tests()
    finalize_reports(static_result, browser_result)
    write_reference_manifest()
    write_package_manifest()
    result = package_zip()
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
